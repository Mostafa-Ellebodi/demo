#################
#### Imports ####
#################
import streamlit as st
import tensorflow_hub as hub
import tensorflow as tf
from textsplit.tools import get_penalty, get_segments
from textsplit.algorithm import split_optimal, split_greedy, get_total
from textsplit.tools import SimpleSentenceTokenizer
import docx
from docx.oxml.shared import OxmlElement # Necessary Import
from docx.oxml import ns
from transformers import  PegasusTokenizerFast, PegasusForConditionalGeneration
import torch
from summa.summarizer import summarize
import os
from pathlib import Path
import glob
from math import ceil
from treelib import Node, Tree
from text_preprocessing import preprocess_text
from text_preprocessing import *
from PIL import Image
import shutil
import gc


###################
#### Functions ####
###################

#helper function for creating new empty directories
def clean_mkdir(path):
    if Path(path).exists():
        shutil.rmtree(path)
    os.makedirs(path)

def upload_s3():
    os.system('aws s3 cp TOC.docx s3://streamlitdocx/TOC.docx')
    os.system('aws s3api put-object-acl --bucket streamlitdocx --key report.docx --acl public-read')
    print("FILE UPLOADED SUCCEFULLY")
    print()
    print()

def cleaning(segment):
    print("entered cleaning function")
    preprocess_functions = [
                            remove_whitespace,
                            remove_special_character, 
                            normalize_unicode, 
                            expand_contraction,
                            remove_name,
                            ]
    
    segment = preprocess_text(segment, preprocess_functions)
    
    return segment

def do_segmentation(input_text, segment_len):
    sentenced_text = sentence_tokenizer(input_text.replace("\n", " "))
    sentence_vectors = embed(sentenced_text).numpy()
    penalty = get_penalty([sentence_vectors], segment_len)
    optimal_segmentation = split_optimal(sentence_vectors, penalty, seg_limit=None)
    segmented_text = get_segments(sentenced_text, optimal_segmentation)
    
    segments = []
    for segment_sentences in segmented_text:
        segment = ' '.join(segment_sentences)
        segment = segment.replace("\n","")
        segments.append(segment)    
        
    return segments

def segment_headlining(segment, 
                      max_length = 512,                                     
                      num_beams=1,
                      repetition_penalty=1, 
                      length_penalty=1.0, 
                      early_stopping=False):
    # text cleaning
    segment = cleaning(segment)
    
    len_segment = len(segment.split(" "))
    if  len_segment > max_length:
        text = summarize(segment, words=max_length)
    else:
        text = segment
    tokenized_text = headline_model_tokenizer.encode(text, 
                                                      return_tensors="pt", 
                                                      max_length=max_length, 
                                                      truncation=True, 
                                                      padding=True, 
                                                      pad_to_max_length=True).to(device)
    


        
    headline_ids = headline_model.generate(tokenized_text,
                                           num_beams=num_beams,
                                           length_penalty=length_penalty, 
                                           early_stopping=early_stopping,
                                           no_repeat_ngram_size=1,
                                           max_length = max_length,
                                           bad_words_ids = [[i] for i in vocabulary if (i not in tokenized_text.cpu().numpy()[0]) and (int(i) < 96103)]
                                           )
    headline = headline_model_tokenizer.batch_decode(headline_ids, skip_special_tokens=True)
    
    return headline

class segment_class(object):
        def __init__(self, segment, headline, num_sentences, num_tokens, level, branch, summary):
            self.segment = segment
            self.headline = headline
            self.num_sentences = num_sentences
            self.num_tokens = num_tokens
            self.level = level
            self.branch = branch
            self.summary = summary

def segment_annotation(segment, branch, num_sentences, level):
    headline = segment_headlining(segment)
    
    
    headline = '.'.join( [str(int(i) + 1) for i in branch[1:]]) + ". " + ' / '.join(headline)
#     headline = post_processing(headline)
    summary = None
    num_tokens = len(segment.split(" "))
    tree.create_node(headline + " " + str(num_sentences), '_'.join(branch), 
                     parent='_'.join(branch[:-1]),
                     data= segment_class(segment, headline, num_sentences, num_tokens, level, branch, summary)
                    )

    headlines.append(headline)
    summaries.append(summary)

def recursively_process(segment, 
                        init_seg_len, 
                        recursive_segmentation_factor,
                        stopping_len,
                        relative=False, 
                        level=0,
                        branch=['0']):
    
    # get number of sentences
    num_sentences = len(sentence_tokenizer(segment.replace("\n", " ")))
    
    # is this input too short? if yes, we would terminate.
    is_shortest = num_sentences < stopping_len*2
    
    # did the segmentation target length become too low? if yes, we would terminate.
    is_over_segmentation = init_seg_len < stopping_len
    terminate = is_shortest or is_over_segmentation
    
    # is this a leaf node? or we are at root node?
    is_leaf = bool(level)
        
    # is this a root or leaf node?
    if is_leaf:
        segment_annotation(segment, branch, num_sentences, level) # summerize and headline
        
    # check termination conditions for the input segment
    if terminate:
#         print("state: terminate" )
        segments.append(segment)
        return
    
    # initialize sub_segments list
    sub_segments = [segment]
    
    # is this input segmentable by the targent length? or is it small?
    if (num_sentences >= init_seg_len*2):
        sub_segments = do_segmentation(segment, segment_len= init_seg_len)
    
    # is this segment small for the given segmentation target length or segmentation returned one segment?
    if (num_sentences < init_seg_len*2) or len(sub_segments) == 1:  # if yes
        
        init_seg_len = int(num_sentences / 2)  # , make the target small enough
        
        # while we did not manage to segement or already reach termination .. 
        while(init_seg_len >= stopping_len  and len(sub_segments) == 1):
            sub_segments = do_segmentation(segment, segment_len= init_seg_len)
            init_seg_len-=1
        init_seg_len+=1
        
    if len(sub_segments) == 1:
        segments.append(segment)
        return
    elif not(is_leaf): 
        pass
    else:
        segments.append("")
            
    if not(relative):
        next_seg_len = int(init_seg_len / recursive_segmentation_factor)
       
    level+=1
    for b, sub_segment in enumerate(sub_segments):
        branch.append(str(b))
        if relative:
            num_sentences_sub = len(sentence_tokenizer(sub_segment.replace("\n", " ")))
            next_seg_len = int(num_sentences_sub / recursive_segmentation_factor)
        recursively_process(segment = sub_segment, 
                            init_seg_len = next_seg_len, 
                            recursive_segmentation_factor = recursive_segmentation_factor, 
                            stopping_len = stopping_len,
                            relative=relative,
                            level=level,
                            branch=branch)
        branch.pop()

def create_element(name):
    return OxmlElement(name)
 
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)
 
 
def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
 
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
 
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
 
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
 
 
def iter_paragraphs(parent, recursive=True):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(repr(type(parent)))
    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
          yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
          if recursive:
            for row in docx.table.Table(child, parent).table.rows:
                for cell in row.cells:
                  for child_paragraph in iter_paragraphs(cell):
                      yield child_paragraph

def build_docx(segments, summaries, headlines):
    
    counter = ceil(len(headlines) / 15) + 1
    log_page_no = []
    mydoc = docx.Document()
    for my_headline, summary, segment in zip(headlines, summaries, segments):
      log_page_no.append(counter)
      tokenized_segment = segment.split(" ")
      L = len(tokenized_segment)
      start_range = 300
      step = 400
      i=0
      mydoc.add_heading(my_headline.strip(), 0)
      mydoc.add_paragraph().add_run(summary).bold = True
      if L >= start_range:
        mydoc.add_paragraph().add_run(" ".join(tokenized_segment[:start_range]))
        mydoc.add_page_break()
        counter += 1
        i = start_range
      for i in range(550, L, step):
        mydoc.add_paragraph().add_run(" ".join(tokenized_segment[i-step:i]))
        mydoc.add_page_break()
        counter += 1
      mydoc.add_paragraph().add_run(" ".join(tokenized_segment[i:]))
      mydoc.add_page_break()
      counter += 1


    mydoc = docx.Document()
    add_page_number(mydoc.sections[0].footer.paragraphs[0].add_run())
    mydoc.add_heading("Table of Contents", 0)
    counter_TOC = 0
    for headline, page_no in zip(headlines, log_page_no):
      counter_TOC+=1
      mydoc.add_paragraph(headline + " .... " + str(page_no))
      if counter_TOC % 15 == 0 and counter_TOC != len(headlines):
        mydoc.add_page_break()
        
    return mydoc

def process_file(TEXT_FILE_PATH, 
                 OUTPUT_FILE_PATH, 
                 init_seg_len=80,
                 recursive_segmentation_factor=2,
                 stopping_len=40,
                 relative=False
                ):
 
    with open(TEXT_FILE_PATH, 'rt', errors='ignore') as f:
      input_text = f.read()
    
    
    
    file_name = os.path.split(TEXT_FILE_PATH)[-1].split(".")[0]
    num_sentences = len(sentence_tokenizer(input_text.replace("\n", " ")))
    
    sentences = sentence_tokenizer(input_text.replace("\n", " ").strip())
    sentences = [",".join(j for j in i.split(',') if len(j) > 2) for i in sentences]
    input_text = ''.join([i for i in sentences if len(i.split(" ")) > 7])
    num_tokens = len(input_text.split(" "))
    tree.create_node(file_name, 
                     '_'.join(['0']),
                     data= segment_class(None, file_name, num_sentences, num_tokens, 0, 0, None) ) # root node
    
    
    recursively_process(input_text, 
                        init_seg_len=init_seg_len, 
                        recursive_segmentation_factor=recursive_segmentation_factor, 
                        stopping_len=stopping_len,
                        relative=relative
                       )
    

    mydoc = build_docx(segments, summaries, headlines)
    file_name = os.path.split(TEXT_FILE_PATH)[-1].split(".")[0]
    mydoc.save(os.path.join(OUTPUT_FILE_PATH, file_name + ".docx"))

# ##############################
# #### TF GPU Configuration ####
# ##############################
# gpus = tf.config.experimental.list_physical_devices('GPU')
# if gpus:
#   try:
#     tf.config.experimental.set_virtual_device_configuration(gpus[0], [tf.config.experimental.VirtualDeviceConfiguration(memory_limit=4096)])
#   except RuntimeError as e:
#     print(e)


###################################
#### Process Files and Download#### #######################
###################################

############################
#### setup basic layout ####
############################

# To disable warnings
st.set_option('deprecation.showfileUploaderEncoding', False)

img = Image.open("tc-logo.png")
st.image(img)

st.write("""
  # Get the best TOC of your text file!
""")
# st.subheader("upload your text file and look at TOC")



st.write("Upload your text files, and pick values from the slide bars on the left. Then, Hit Summarize")
uploaded_files  = st.file_uploader("Please upload your text file: ", type ="txt", accept_multiple_files=True)

if uploaded_files==[]:
  st.write("Please Upload Your files")
  st.stop()


#################
#### Configs ####
#################
INITIAL_SEGMENTATION_LENGTH = st.sidebar.slider('Inital segmentation length (Recommended: 50)', 10, 80)
RECURSIVE_SEGMENTATION_FACTOR = st.sidebar.slider('Recursive Segmentation Factor. (Recommended: 2)', 1, 5)
STOPPING_LENGTH = st.sidebar.slider('Stopping Length. (Recommended: 50)', 10, 80)
_RELATIVE = st.sidebar.selectbox(
    'Do you want the segmentatio length to be realtive?',
    ('True', 'False')
)
if _RELATIVE == 'True':
  RELATIVE = True
elif _RELATIVE == 'False':
  RELATIVE = False

OUTPUT_PATH = "./output"
clean_mkdir(OUTPUT_PATH)
DATA_PATH = "./input"
clean_mkdir(DATA_PATH)
class model_args:
    DATA_PATH = DATA_PATH
    OUTPUT_PATH = OUTPUT_PATH
    model_name_or_path = 'pegasus_b8_lr-5_NoEarly_clean_large_AllNews_4/checkpoint-9424'
    INITIAL_SEGMENTATION_LENGTH = INITIAL_SEGMENTATION_LENGTH # the target number of sentences to be contained in the first segmentation iteration
    RECURSIVE_SEGMENTATION_FACTOR = RECURSIVE_SEGMENTATION_FACTOR # the factor at which each segment is furthure split
    STOPPING_LENGTH = STOPPING_LENGTH # the length at which segmentation should stop
    RELATIVE = RELATIVE # whether the target length at each segmentation iteration should be factor split based on previous target legth "RELATIVE=False" or relative to previous segement length "RELATIVE=TRUE"

##############################
#### Write uploaded files ####
##############################
print("writing uploaded files")
for uploaded_file in uploaded_files:
    input_text = uploaded_file.read().decode()
    fn = uploaded_file.name
    filePath = DATA_PATH + fn
    # Write-Overwrites 
    file1 = open(filePath,"w")#write mode 
    file1.write(input_text) 
    file1.close() 


###################################
#### Process Files and Download####
###################################
if st.button("Get TOC"):
    try:
        #########################
        #### Load all models ####
        #########################
        print("start loading models")
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        USE_path = "USEV4/"
        embed = hub.load(USE_path)

        sentence_tokenizer = SimpleSentenceTokenizer()

        headline_model_tokenizer = PegasusTokenizerFast.from_pretrained(model_args.model_name_or_path)
        vocabulary = list(headline_model_tokenizer.get_vocab().values())

        headline_model = PegasusForConditionalGeneration.from_pretrained(model_args.model_name_or_path)
        print("end loading models")


        docs_trees = []
        docs_segments = []
        docs_summaries = []
        docs_headlines = []

        paths = glob.glob(os.path.join(model_args.DATA_PATH, "*.txt"))
        for path in paths:
            print("processing first file")
            segments = []
            summaries = []
            headlines = []
            tree = Tree()
            print("second step file")

            process_file(path, 
                        model_args.OUTPUT_PATH, 
                        model_args.INITIAL_SEGMENTATION_LENGTH, 
                        model_args.RECURSIVE_SEGMENTATION_FACTOR,
                        model_args.STOPPING_LENGTH,
                        model_args.RELATIVE
                        )
            print("finished processing step file")
    #     tree.show()

    #     docs_trees.append(tree)
    #     docs_segments.append(segments)
    #     docs_summaries.append(summaries)
    #     docs_headlines.append(headlines)

        # mydoc.save("report.docx")

        # upload_s3()
        # link = "https://streamlitdocx.s3.us-east-2.amazonaws.com/report.docx"
        # st.markdown(f"[Download Report.docx]({link})", unsafe_allow_html = True)
        print("start cleaning")
        gc.collect()
        st.caching.clear_cache()
        del headline_model_tokenizer, headline_model
        file = None
        if not file:
            st.write("Process finished. You can upload new files")
            # st.stop()
        print("end cleaning")
    except:
        st.write("ERROR: Reload app and make sure you are entering values correctly and reupload file")